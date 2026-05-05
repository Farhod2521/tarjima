"""
DOCX document formatting endpoints.

The endpoint polishes paragraph text after translation or OCR and applies
uniform font settings to the resulting DOCX file.
"""
import asyncio
import json
import logging
import uuid
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.shared import Pt
from fastapi import APIRouter, BackgroundTasks, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse, StreamingResponse
from openai import AsyncOpenAI

from config import settings

router = APIRouter(prefix="/document-format", tags=["document-format"])
logger = logging.getLogger(__name__)

FORMAT_DIR = Path("/data/document_format")
FORMAT_DIR.mkdir(parents=True, exist_ok=True)

_jobs: dict[str, dict[str, Any]] = {}

LANG_PROMPT = {
    "uzbek": "Uzbek, Latin script, modern official style",
    "russian": "Russian",
    "english": "English",
}

SUPPORTED_FONTS = {
    "Times New Roman",
    "Arial",
    "Calibri",
    "Cambria",
    "Georgia",
    "Aptos",
}


def _iter_paragraphs(doc: DocxDocument):
    for para in doc.paragraphs:
        if para.text.strip():
            yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        yield para


def _apply_font(para, font_family: str, font_size: int) -> None:
    if not para.runs:
        para.add_run("")
    for run in para.runs:
        run.font.name = font_family
        run.font.size = Pt(font_size)


def _replace_text_and_style(para, new_text: str, font_family: str, font_size: int) -> None:
    if not para.runs:
        run = para.add_run(new_text)
        run.font.name = font_family
        run.font.size = Pt(font_size)
        return

    para.runs[0].text = new_text
    para.runs[0].font.name = font_family
    para.runs[0].font.size = Pt(font_size)

    for run in para.runs[1:]:
        run.text = ""

    _apply_font(para, font_family, font_size)


def _supports_reasoning(model: str) -> bool:
    return model.startswith("gpt-5")


def _reasoning_kwargs(model: str) -> dict[str, Any]:
    effort = settings.openai_reasoning_effort.strip().lower()
    if not effort or not _supports_reasoning(model):
        return {}
    return {"reasoning": {"effort": effort}}


def _response_text(response: Any) -> str:
    text = getattr(response, "output_text", None)
    if isinstance(text, str):
        return text

    parts: list[str] = []
    for item in getattr(response, "output", []) or []:
        for content in getattr(item, "content", []) or []:
            content_text = getattr(content, "text", None)
            if isinstance(content_text, str):
                parts.append(content_text)
    return "".join(parts)


def _format_system_prompt(document_lang: str) -> str:
    return (
        "You are a meticulous editor for official DOCX documents. "
        f"The paragraph language is {LANG_PROMPT[document_lang]}. "
        "Polish the paragraph without translating it into another language. "
        "Return only the corrected paragraph text and nothing else. "
        "Fix punctuation, commas, periods, colons, semicolons, quotation marks, "
        "parentheses, repeated spaces, broken line-break artifacts, accidental "
        "spacing around punctuation, inconsistent capitalization, and obvious "
        "post-translation formatting noise. "
        "Join suffixes or particles back to the correct word only when they were "
        "clearly split by mistake. "
        "For Uzbek output, use natural modern Uzbek Latin and ASCII apostrophes "
        "such as o' and g'. "
        "Preserve meaning, sentence order, numbers, dates, references, list marks, "
        "document codes, standards, abbreviations, names, units, and legal or "
        "technical terminology. "
        "Do not add commentary, explanations, markdown, quotes, or alternative "
        "versions. "
        "If the paragraph is already correct, return it unchanged."
    )


async def _format_text(
    text: str,
    document_lang: str,
    client: AsyncOpenAI,
) -> str:
    if not text.strip():
        return text

    model = settings.translation_model
    messages = [
        {"role": "system", "content": _format_system_prompt(document_lang)},
        {
            "role": "user",
            "content": (
                "Correct the paragraph and keep the same language and meaning.\n\n"
                f"{text}"
            ),
        },
    ]

    try:
        try:
            response = await client.responses.create(
                model=model,
                input=messages,
                max_output_tokens=4096,
                **_reasoning_kwargs(model),
            )
            result = _response_text(response).strip()
        except AttributeError:
            logger.warning("OpenAI SDK has no Responses API; falling back to Chat Completions.")
            response = await client.chat.completions.create(
                model=model,
                messages=messages,
                max_completion_tokens=4096,
            )
            result = (response.choices[0].message.content or "").strip()

        return result if result else text
    except Exception as exc:
        logger.exception("Document formatting failed: %s", exc)
        return text


async def _do_document_formatting(
    job_id: str,
    input_path: Path,
    output_path: Path,
    document_lang: str,
    font_family: str,
    font_size: int,
) -> None:
    job = _jobs[job_id]
    try:
        doc = DocxDocument(str(input_path))
        paragraphs = list(_iter_paragraphs(doc))
        total = len(paragraphs)
        job.update({"total": total, "status": "processing"})

        if total == 0:
            doc.save(str(output_path))
            job.update({"status": "done", "progress": 100})
            return

        client = AsyncOpenAI(
            api_key=settings.openai_api_key,
            timeout=settings.openai_timeout_seconds,
        )
        semaphore = asyncio.Semaphore(settings.max_parallel_translations)

        async def process(para) -> None:
            async with semaphore:
                corrected = await _format_text(para.text, document_lang, client)
                _replace_text_and_style(para, corrected, font_family, font_size)
                job["done"] += 1
                job["progress"] = int(job["done"] / total * 100)

        await asyncio.gather(*[process(para) for para in paragraphs])
        doc.save(str(output_path))
        job.update({"status": "done", "progress": 100})

    except Exception as exc:
        logger.exception("Document formatting job %s failed: %s", job_id, exc)
        job.update({"status": "error", "error": str(exc)})
    finally:
        input_path.unlink(missing_ok=True)


@router.post("/")
async def start_document_formatting(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    document_lang: str = Form(...),
    font_family: str = Form(...),
    font_size: int = Form(...),
):
    if not settings.openai_api_key:
        raise HTTPException(500, "OPENAI_API_KEY sozlanmagan")
    if document_lang not in LANG_PROMPT:
        raise HTTPException(400, "Noma'lum hujjat tili")
    if font_family not in SUPPORTED_FONTS:
        raise HTTPException(400, "Tanlangan shrift qo'llab-quvvatlanmaydi")
    if font_size < 8 or font_size > 32:
        raise HTTPException(400, "Shrift o'lchami 8 va 32 oralig'ida bo'lishi kerak")

    filename = file.filename or "document.docx"
    if not filename.lower().endswith(".docx"):
        raise HTTPException(400, "Faqat .docx fayl qabul qilinadi")

    content = await file.read()
    max_bytes = settings.max_file_size_mb * 1024 * 1024
    if len(content) > max_bytes:
        raise HTTPException(400, f"Fayl {settings.max_file_size_mb}MB dan oshmasligi kerak")

    job_id = str(uuid.uuid4())
    input_path = FORMAT_DIR / f"{job_id}_in.docx"
    output_path = FORMAT_DIR / f"{job_id}_out.docx"
    input_path.write_bytes(content)

    _jobs[job_id] = {
        "status": "pending",
        "progress": 0,
        "total": 0,
        "done": 0,
        "error": None,
        "filename": filename,
        "document_lang": document_lang,
        "font_family": font_family,
        "font_size": font_size,
        "model": settings.translation_model,
    }

    background_tasks.add_task(
        _do_document_formatting,
        job_id,
        input_path,
        output_path,
        document_lang,
        font_family,
        font_size,
    )
    return {"job_id": job_id}


@router.get("/{job_id}/progress")
async def document_format_progress(job_id: str):
    if job_id not in _jobs:
        raise HTTPException(404, "Jarayon topilmadi")

    async def stream():
        while True:
            job = _jobs.get(job_id, {})
            yield f"data: {json.dumps(job)}\n\n"
            if job.get("status") in ("done", "error"):
                break
            await asyncio.sleep(0.4)

    return StreamingResponse(
        stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@router.get("/{job_id}/download")
def download_formatted_document(job_id: str):
    job = _jobs.get(job_id)
    if not job:
        raise HTTPException(404, "Jarayon topilmadi")
    if job["status"] != "done":
        raise HTTPException(400, "Hujjatni tartiblash hali tugamagan")

    output_path = FORMAT_DIR / f"{job_id}_out.docx"
    if not output_path.exists():
        raise HTTPException(404, "Fayl topilmadi")

    original_name = job.get("filename", "document.docx")
    return FileResponse(
        path=str(output_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"tartiblangan_{original_name}",
    )
