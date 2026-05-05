"""
DOCX translation endpoints.

Only text is translated; images, tables, and the general DOCX structure stay in place.
"""
import asyncio
import json
import logging
import uuid
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from fastapi import APIRouter, BackgroundTasks, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse, StreamingResponse
from openai import AsyncOpenAI

from config import settings

router = APIRouter(prefix="/translate", tags=["translate"])
logger = logging.getLogger(__name__)

TRANS_DIR = Path("/data/translations")
TRANS_DIR.mkdir(parents=True, exist_ok=True)

_jobs: dict[str, dict[str, Any]] = {}

LANG_PROMPT = {
    "uzbek": "Uzbek, Latin script, modern official style",
    "russian": "Russian",
    "english": "English",
}


def _iter_paragraphs(doc: DocxDocument):
    for para in doc.paragraphs:
        if para.text.strip():
            yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        yield para


def _replace_text(para, new_text: str) -> None:
    runs = para.runs
    if not runs:
        return
    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""


def _supports_reasoning(model: str) -> bool:
    return model.startswith("gpt-5")


def _reasoning_kwargs(model: str) -> dict[str, Any]:
    effort = settings.openai_reasoning_effort.strip().lower()
    if not effort or not _supports_reasoning(model):
        return {}
    return {"reasoning": {"effort": effort}}


def _response_text(response: Any) -> str:
    text = getattr(response, "output_text", None)
    if isinstance(text, str):
        return text

    parts: list[str] = []
    for item in getattr(response, "output", []) or []:
        for content in getattr(item, "content", []) or []:
            content_text = getattr(content, "text", None)
            if isinstance(content_text, str):
                parts.append(content_text)
    return "".join(parts)


def _translation_system_prompt(from_lang: str, to_lang: str) -> str:
    return (
        "You are a professional translator for technical and regulatory documents. "
        f"Translate from {LANG_PROMPT[from_lang]} to {LANG_PROMPT[to_lang]}. "
        "Return only the translated text. "
        "Preserve numbers, formulas, units, tables' cell meaning, standards, document codes, "
        "abbreviations, product names, organization names, and proper nouns. "
        "Do not add explanations, comments, quotes, markdown, or the original text. "
        "Keep paragraph meaning and tone. "
        "For Uzbek output, use natural modern Uzbek Latin and ASCII apostrophes such as o' and g'."
    )


async def _translate_text(
    text: str,
    from_lang: str,
    to_lang: str,
    client: AsyncOpenAI,
) -> str:
    if not text.strip():
        return text

    model = settings.translation_model
    messages = [
        {"role": "system", "content": _translation_system_prompt(from_lang, to_lang)},
        {"role": "user", "content": text},
    ]

    try:
        try:
            response = await client.responses.create(
                model=model,
                input=messages,
                max_output_tokens=4096,
                **_reasoning_kwargs(model),
            )
            result = _response_text(response).strip()
        except AttributeError:
            logger.warning("OpenAI SDK has no Responses API; falling back to Chat Completions.")
            response = await client.chat.completions.create(
                model=model,
                messages=messages,
                max_completion_tokens=4096,
            )
            result = (response.choices[0].message.content or "").strip()

        return result if result else text
    except Exception as exc:
        logger.exception("Translation failed: %s", exc)
        return text


async def _do_translation(
    job_id: str,
    input_path: Path,
    output_path: Path,
    from_lang: str,
    to_lang: str,
) -> None:
    job = _jobs[job_id]
    try:
        doc = DocxDocument(str(input_path))
        paragraphs = list(_iter_paragraphs(doc))
        total = len(paragraphs)
        job.update({"total": total, "status": "processing"})

        if total == 0:
            doc.save(str(output_path))
            job.update({"status": "done", "progress": 100})
            return

        client = AsyncOpenAI(
            api_key=settings.openai_api_key,
            timeout=settings.openai_timeout_seconds,
        )
        semaphore = asyncio.Semaphore(settings.max_parallel_translations)

        async def process(para) -> None:
            async with semaphore:
                original = para.text
                translated = await _translate_text(original, from_lang, to_lang, client)
                _replace_text(para, translated)
                job["done"] += 1
                job["progress"] = int(job["done"] / total * 100)

        await asyncio.gather(*[process(p) for p in paragraphs])
        doc.save(str(output_path))
        job.update({"status": "done", "progress": 100})

    except Exception as exc:
        logger.exception("Translation job %s failed: %s", job_id, exc)
        job.update({"status": "error", "error": str(exc)})
    finally:
        input_path.unlink(missing_ok=True)


@router.post("/")
async def start_translation(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    from_lang: str = Form(...),
    to_lang: str = Form(...),
):
    if not settings.openai_api_key:
        raise HTTPException(500, "OPENAI_API_KEY sozlanmagan")
    if from_lang not in LANG_PROMPT or to_lang not in LANG_PROMPT:
        raise HTTPException(400, "Noma'lum til")
    if from_lang == to_lang:
        raise HTTPException(400, "Manba va maqsad tili bir xil bo'lmasligi kerak")

    filename = file.filename or "document.docx"
    if not filename.lower().endswith(".docx"):
        raise HTTPException(400, "Faqat .docx fayl qabul qilinadi")

    content = await file.read()
    max_bytes = settings.max_file_size_mb * 1024 * 1024
    if len(content) > max_bytes:
        raise HTTPException(400, f"Fayl {settings.max_file_size_mb}MB dan oshmasligi kerak")

    job_id = str(uuid.uuid4())
    input_path = TRANS_DIR / f"{job_id}_in.docx"
    output_path = TRANS_DIR / f"{job_id}_out.docx"
    input_path.write_bytes(content)

    _jobs[job_id] = {
        "status": "pending",
        "progress": 0,
        "total": 0,
        "done": 0,
        "error": None,
        "filename": filename,
        "to_lang": to_lang,
        "model": settings.translation_model,
    }

    background_tasks.add_task(
        _do_translation,
        job_id,
        input_path,
        output_path,
        from_lang,
        to_lang,
    )
    return {"job_id": job_id}


@router.get("/{job_id}/progress")
async def job_progress(job_id: str):
    if job_id not in _jobs:
        raise HTTPException(404, "Jarayon topilmadi")

    async def stream():
        while True:
            job = _jobs.get(job_id, {})
            yield f"data: {json.dumps(job)}\n\n"
            if job.get("status") in ("done", "error"):
                break
            await asyncio.sleep(0.4)

    return StreamingResponse(
        stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@router.get("/{job_id}/download")
def download_translated(job_id: str):
    job = _jobs.get(job_id)
    if not job:
        raise HTTPException(404, "Jarayon topilmadi")
    if job["status"] != "done":
        raise HTTPException(400, "Tarjima hali tugamagan")

    output_path = TRANS_DIR / f"{job_id}_out.docx"
    if not output_path.exists():
        raise HTTPException(404, "Fayl topilmadi")

    original_name = job.get("filename", "document.docx")
    return FileResponse(
        path=str(output_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"tarjima_{original_name}",
    )
