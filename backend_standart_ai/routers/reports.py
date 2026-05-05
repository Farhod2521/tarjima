"""
Xodimlar hisobati — yaratish, ro'yxat, yuklab olish.
Har bir xodim uchun BITTA Word fayl. Yangi tahlil qo'shilganda qayta generatsiya.
"""
import datetime
import json
import re
import unicodedata
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from database import REPORTS_DIR, Report, get_db

router = APIRouter(prefix="/reports", tags=["reports"])

# ── Rang palitasi ─────────────────────────────────────────────────────────────
PRIMARY   = RGBColor(0x1A, 0x56, 0x76)   # To'q ko'k
SECONDARY = RGBColor(0x2E, 0x9C, 0x9C)   # Teal
ACCENT    = RGBColor(0xF0, 0x78, 0x28)   # To'q sariq-to'q
LIGHT_BG  = RGBColor(0xF0, 0xF7, 0xFA)   # Och ko'k fon
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT = RGBColor(0x1A, 0x1A, 0x2E)
GRAY_TEXT = RGBColor(0x55, 0x65, 0x75)
ROW_ALT   = RGBColor(0xF8, 0xFC, 0xFD)   # Alternativ qator


# ── Schemas ───────────────────────────────────────────────────────────────────

class IssueItem(BaseModel):
    original: str
    corrected: str
    page: Optional[int] = None


class ReportCreate(BaseModel):
    employee_name: str
    standard_name: str
    issues: list[IssueItem]


class ReportOut(BaseModel):
    id: int
    employee_name: str
    standard_name: str
    docx_filename: str
    created_at: datetime.datetime
    standards_count: int = 1

    model_config = {"from_attributes": True}


# ── Yordamchi funksiyalar ─────────────────────────────────────────────────────

def _slug(name: str) -> str:
    """Xodim ismidan fayl nomi yasaydi."""
    name = unicodedata.normalize("NFKD", name).encode("ascii", "ignore").decode("ascii")
    name = re.sub(r"[^\w\s-]", "", name, flags=re.ASCII).strip()
    slug = re.sub(r"[\s]+", "_", name).lower()[:60]
    return slug or "xodim"


def _report_filename(employee_name: str, report_id: int) -> str:
    """Har bir ko'rinadigan report qatori uchun collision bo'lmaydigan fayl nomi."""
    return f"{_slug(employee_name)}_{report_id}.docx"


def _report_path(filename: str) -> Path:
    """Fayl nomini /data/reports ichida saqlashga majburlaydi."""
    return REPORTS_DIR / Path(filename).name


def _employee_report_rows(db: Session, employee_name: str) -> list[dict]:
    records = (
        db.query(Report)
        .filter(Report.employee_name == employee_name)
        .order_by(Report.id.asc())
        .all()
    )
    rows: list[dict] = []
    for rec in records:
        issues = json.loads(rec.issues_json) if rec.issues_json else []
        rows.append({
            "standard_name": rec.standard_name,
            "issues": issues,
            "created_at": rec.created_at,
        })
    return rows


def _delete_report_files(records: list[Report]) -> None:
    filenames: set[str] = set()
    for rec in records:
        if rec.docx_filename:
            filenames.add(rec.docx_filename)
        filenames.add(_report_filename(rec.employee_name, rec.id))

    for filename in filenames:
        _report_path(filename).unlink(missing_ok=True)


def _set_cell_bg(cell, color: RGBColor) -> None:
    """Jadval katakchasi fonini o'rnatadi."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell, color_hex: str = "C5D9E0") -> None:
    """Yupqa chegara."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color_hex)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def _cell_text(cell, text: str, bold=False, color: RGBColor = DARK_TEXT,
               size_pt: int = 10, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color


def _add_page_break(doc) -> None:
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_break_type())


def docx_break_type():
    from docx.enum.text import WD_BREAK
    return WD_BREAK.PAGE


def _set_col_width(table, col_idx: int, width_cm: float) -> None:
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


# ── DOCX generatsiya ──────────────────────────────────────────────────────────

def _generate_employee_docx(
    employee_name: str,
    all_reports: list[dict],   # [{"standard_name": ..., "issues": [...], "created_at": ...}]
    path: Path,
) -> None:
    doc = DocxDocument()

    # Sahifa o'lchamlari: A4
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    # ── BIRINCHI SAHIFA: Sarlavha ─────────────────────────────────────────────

    # Logo / brend sarlavhasi
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_p.add_run("StandartAI")
    logo_run.bold = True
    logo_run.font.size = Pt(28)
    logo_run.font.color.rgb = PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("GRAMMATIK TAHLIL TIZIMI")
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = SECONDARY
    sub_run.font.all_caps = True

    doc.add_paragraph()

    # Gorizontal chiziq (bitta katakchali jadval orqali)
    divider = doc.add_table(rows=1, cols=1)
    divider.rows[0].cells[0].text = ""
    _set_cell_bg(divider.rows[0].cells[0], SECONDARY)
    divider.rows[0].height = Cm(0.05)
    divider.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(0)
    divider.rows[0].cells[0].paragraphs[0].paragraph_format.space_after  = Pt(0)

    doc.add_paragraph()

    # Hisobot sarlavhasi
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("XODIM GRAMMATIK TAHLIL HISOBOTI")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = DARK_TEXT

    doc.add_paragraph()

    # Xodim ismi kartochkasi
    card_tbl = doc.add_table(rows=2, cols=2)
    labels = ["Xodim:", "Sana:"]
    values = [
        employee_name,
        datetime.datetime.now().strftime("%d.%m.%Y"),
    ]
    for i, (lbl, val) in enumerate(zip(labels, values)):
        lc, vc = card_tbl.rows[i].cells[0], card_tbl.rows[i].cells[1]
        _set_cell_bg(lc, LIGHT_BG)
        _set_cell_bg(vc, WHITE)
        _set_cell_borders(lc, "C5D9E0")
        _set_cell_borders(vc, "C5D9E0")
        _cell_text(lc, lbl, bold=True, color=PRIMARY, size_pt=10)
        _cell_text(vc, val, color=DARK_TEXT, size_pt=10)
        lc.width = Cm(4)
        vc.width = Cm(13)

    doc.add_paragraph()

    # ── STATISTIKA JADVALI ────────────────────────────────────────────────────
    stat_p = doc.add_paragraph()
    stat_run = stat_p.add_run("Tahlil Statistikasi")
    stat_run.bold = True
    stat_run.font.size = Pt(12)
    stat_run.font.color.rgb = PRIMARY

    stat_tbl = doc.add_table(rows=1, cols=3)
    # Sarlavha qatori
    hdr_cells = stat_tbl.rows[0].cells
    for cell, lbl in zip(hdr_cells, ["Standart nomi", "Xato soni", "Sana"]):
        _set_cell_bg(cell, PRIMARY)
        _set_cell_borders(cell, "1A5676")
        _cell_text(cell, lbl, bold=True, color=WHITE, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    total_errors = 0
    for i, rep in enumerate(all_reports):
        row = stat_tbl.add_row().cells
        bg = ROW_ALT if i % 2 == 0 else WHITE
        for cell in row:
            _set_cell_bg(cell, bg)
            _set_cell_borders(cell, "C5D9E0")
        err_count = len(rep["issues"])
        total_errors += err_count
        _cell_text(row[0], rep["standard_name"], color=DARK_TEXT, size_pt=10)
        _cell_text(row[1], str(err_count), color=ACCENT if err_count > 0 else SECONDARY,
                   bold=True, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        dt = rep.get("created_at", "")
        if isinstance(dt, datetime.datetime):
            dt = dt.strftime("%d.%m.%Y")
        _cell_text(row[2], str(dt), color=GRAY_TEXT, size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Jami qatori
    total_row = stat_tbl.add_row().cells
    for cell in total_row:
        _set_cell_bg(cell, LIGHT_BG)
        _set_cell_borders(cell, "2E9C9C")
    _cell_text(total_row[0], "JAMI", bold=True, color=PRIMARY, size_pt=10)
    _cell_text(total_row[1], str(total_errors), bold=True, color=ACCENT, size_pt=11,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(total_row[2], "", color=GRAY_TEXT, size_pt=9)

    _set_col_width(stat_tbl, 0, 10)
    _set_col_width(stat_tbl, 1, 3)
    _set_col_width(stat_tbl, 2, 4)

    # ── HAR BIR STANDART UCHUN ALOHIDA BO'LIM ────────────────────────────────
    for rep in all_reports:
        # Yangi sahifadan boshlash
        doc.add_page_break()

        # Bo'lim sarlavhasi
        sec_p = doc.add_paragraph()
        sec_run = sec_p.add_run(f"Standart: {rep['standard_name']}")
        sec_run.bold = True
        sec_run.font.size = Pt(13)
        sec_run.font.color.rgb = PRIMARY

        date_p = doc.add_paragraph()
        dt = rep.get("created_at", "")
        if isinstance(dt, datetime.datetime):
            dt = dt.strftime("%d.%m.%Y %H:%M")
        date_run = date_p.add_run(f"Tahlil sanasi: {dt}")
        date_run.font.size = Pt(9)
        date_run.font.color.rgb = GRAY_TEXT
        date_run.italic = True

        doc.add_paragraph()

        issues = rep["issues"]
        if not issues:
            ok_p = doc.add_paragraph()
            ok_run = ok_p.add_run("✓  Xato topilmadi — hujjat to'g'ri yozilgan.")
            ok_run.font.size = Pt(11)
            ok_run.font.color.rgb = SECONDARY
            ok_run.italic = True
            continue

        # Xatolar jadvali
        err_tbl = doc.add_table(rows=1, cols=3)

        # Sarlavha
        h_cells = err_tbl.rows[0].cells
        for cell, lbl in zip(h_cells, ["Xato", "To\u02BBg\u02BBrisi", "Sahifa"]):
            _set_cell_bg(cell, SECONDARY)
            _set_cell_borders(cell, "2E9C9C")
            _cell_text(cell, lbl, bold=True, color=WHITE, size_pt=10,
                       align=WD_ALIGN_PARAGRAPH.CENTER)

        for i, item in enumerate(issues):
            row = err_tbl.add_row().cells
            bg = ROW_ALT if i % 2 == 0 else WHITE
            for cell in row:
                _set_cell_bg(cell, bg)
                _set_cell_borders(cell, "C5D9E0")
            _cell_text(row[0], item.get("original", ""), color=RGBColor(0xC0, 0x20, 0x20), size_pt=10)
            _cell_text(row[1], item.get("corrected", ""), color=RGBColor(0x15, 0x6B, 0x35), size_pt=10)
            page_val = item.get("page")
            page_str = str(page_val) if page_val else "\u2014"
            _cell_text(row[2], page_str, color=GRAY_TEXT, size_pt=10,
                       align=WD_ALIGN_PARAGRAPH.CENTER)

        _set_col_width(err_tbl, 0, 7)
        _set_col_width(err_tbl, 1, 7)
        _set_col_width(err_tbl, 2, 3)

    # Pastki imzo
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("StandartAI — Grammatik Tahlil Tizimi")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = GRAY_TEXT
    footer_run.italic = True

    doc.save(str(path))


# ── Endpoints ─────────────────────────────────────────────────────────────────

@router.get("/", response_model=list[ReportOut])
def list_reports(db: Session = Depends(get_db)):
    # Har bir xodim uchun ENG SO'NGGI yozuvni, standartlar sonini qaytaradi
    all_reports = db.query(Report).order_by(Report.id.desc()).all()
    seen: dict[str, Report] = {}
    counts: dict[str, int] = {}
    for r in all_reports:
        counts[r.employee_name] = counts.get(r.employee_name, 0) + 1
        if r.employee_name not in seen:
            seen[r.employee_name] = r
    result = []
    for emp, rep in seen.items():
        out = ReportOut(
            id=rep.id,
            employee_name=rep.employee_name,
            standard_name=rep.standard_name,
            docx_filename=rep.docx_filename,
            created_at=rep.created_at,
            standards_count=counts[emp],
        )
        result.append(out)
    return result


@router.post("/", response_model=ReportOut, status_code=201)
def create_report(data: ReportCreate, db: Session = Depends(get_db)):
    # Yangi yozuv saqlash
    issues_json = json.dumps(
        [{"original": i.original, "corrected": i.corrected, "page": i.page}
         for i in data.issues],
        ensure_ascii=False,
    )
    report = Report(
        employee_name=data.employee_name,
        standard_name=data.standard_name,
        docx_filename="pending.docx",
        issues_json=issues_json,
    )
    db.add(report)
    db.flush()
    report.docx_filename = _report_filename(data.employee_name, report.id)
    db.commit()
    db.refresh(report)

    # Xodimning BARCHA yozuvlarini yig'ib, DOCX qayta generatsiya
    all_rep_data = _employee_report_rows(db, data.employee_name)
    path = _report_path(report.docx_filename)
    _generate_employee_docx(data.employee_name, all_rep_data, path)

    return report


@router.get("/{report_id}/download")
def download_report(report_id: int, db: Session = Depends(get_db)):
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Hisobot topilmadi")

    # Eski yoki boshqa xodimga tegishli fayl qaytmasligi uchun har safar qayta generatsiya qilamiz.
    filename = _report_filename(report.employee_name, report.id)
    path = _report_path(filename)
    all_rep_data = _employee_report_rows(db, report.employee_name)
    _generate_employee_docx(report.employee_name, all_rep_data, path)
    if report.docx_filename != filename:
        if report.docx_filename:
            _report_path(report.docx_filename).unlink(missing_ok=True)
        report.docx_filename = filename
        db.commit()

    safe_name = f"{report.employee_name}_hisobot.docx".replace("/", "_")
    return FileResponse(
        path=str(path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=safe_name,
        headers={"Cache-Control": "no-store"},
    )


@router.delete("/{report_id}", status_code=204)
def delete_report(report_id: int, db: Session = Depends(get_db)):
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Hisobot topilmadi")

    # Hisobot sahifasida bitta qator bitta xodimni bildiradi.
    # Shuning uchun o'chirish shu xodimning barcha report yozuvlarini tozalaydi.
    all_records = (
        db.query(Report)
        .filter(Report.employee_name == report.employee_name)
        .all()
    )
    _delete_report_files(all_records)

    for rec in all_records:
        db.delete(rec)
    db.commit()
