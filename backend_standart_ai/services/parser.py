"""
Fayl parser servisi.
PDF: PyMuPDF (fitz) — tez va ishonchli
DOCX: python-docx — paragraf tartibini saqlaydi
"""
from dataclasses import dataclass, field
from pathlib import Path
import logging

import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)


@dataclass
class Paragraph:
    index: int
    text: str
    page_number: int = 0   # PDF uchun qaysi sahifada


@dataclass
class ParseResult:
    paragraphs: list[Paragraph] = field(default_factory=list)
    page_count: int = 0
    warnings: list[str] = field(default_factory=list)


def parse_pdf(path: Path) -> ParseResult:
    """
    PDF fayldan matnni sahifa-sahifa o'qib, paragraflar ro'yxatini qaytaradi.
    Skanerlangan (image-only) sahifalar uchun warning chiqaradi.
    """
    result = ParseResult()
    para_index = 0

    doc = fitz.open(str(path))
    result.page_count = len(doc)

    for page_num, page in enumerate(doc, start=1):
        text = page.get_text("text").strip()

        if not text:
            result.warnings.append(
                f"{page_num}-sahifada matn topilmadi — bu sahifa skanerlangan bo'lishi mumkin (OCR kerak)"
            )
            continue

        # Har bir bo'sh qatorni paragraf chegarasi deb olamiz
        raw_paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

        # Agar \n\n bo'linish ko'p paragraf bermasa, \n bilan sinab ko'r
        if len(raw_paragraphs) <= 1:
            raw_paragraphs = [p.strip() for p in text.split("\n") if p.strip()]

        for raw in raw_paragraphs:
            if len(raw) > 10:  # juda qisqa qatorlarni o'tkazib yubor
                result.paragraphs.append(Paragraph(index=para_index, text=raw, page_number=page_num))
                para_index += 1

    doc.close()

    if not result.paragraphs:
        result.warnings.append("Faylda o'qib bo'ladigan matn topilmadi")

    logger.info("PDF parsed: %d pages, %d paragraphs, %d warnings",
                result.page_count, len(result.paragraphs), len(result.warnings))
    return result


def parse_docx(path: Path) -> ParseResult:
    """
    DOCX fayldan paragraflarni tartib bilan o'qiydi.
    Sahifa raqami matn hajmiga asoslan taxminiy hisoblanadi (~2500 belgi = 1 sahifa).
    """
    result = ParseResult()

    doc = Document(str(path))
    para_index = 0
    cumulative_chars = 0
    CHARS_PER_PAGE = 2500  # A4, 12pt, taxminiy

    def append_paragraph(text: str) -> None:
        nonlocal para_index, cumulative_chars
        text = text.strip()
        if text and len(text) > 5:
            estimated_page = max(1, cumulative_chars // CHARS_PER_PAGE + 1)
            result.paragraphs.append(Paragraph(index=para_index, text=text, page_number=estimated_page))
            cumulative_chars += len(text) + 1
            para_index += 1

    for para in doc.paragraphs:
        append_paragraph(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    append_paragraph(para.text)

    result.page_count = max(1, cumulative_chars // CHARS_PER_PAGE + 1) if para_index > 0 else 0

    if not result.paragraphs:
        result.warnings.append("DOCX faylda matn topilmadi")

    logger.info("DOCX parsed: %d paragraphs, ~%d pages", len(result.paragraphs), result.page_count)
    return result
