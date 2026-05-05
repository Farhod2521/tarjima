import logging
import tempfile
import time
from pathlib import Path

from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse

from config import settings
from database import init_db
from routers.auth import router as auth_router
from routers.document_format import router as document_format_router
from routers.employees import router as employees_router
from routers.history import router as history_router
from routers.reports import router as reports_router
from routers.translator import router as translator_router
from routers.users import router as users_router
from schemas import AnalysisResponse, FileMetadata, Summary
from services.analyzer import analyze_all_chunks, build_summary
from services.chunker import build_chunks
from services.parser import parse_docx, parse_pdf

# ── Fayl saqlash uchun temp dir ──────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph as RLParagraph, Spacer
from docx import Document as DocxDocument
import uuid
from typing import Literal

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(name)s | %(message)s",
)
logger = logging.getLogger(__name__)

UPLOAD_DIR = Path(tempfile.gettempdir()) / "standart_tahlil"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# ── App ───────────────────────────────────────────────────────────────────────

app = FastAPI(
    title="Standart Tahlil AI",
    description="Hujjatlardagi aniq korrektura xatolarini tahlil qiluvchi API",
    version="2.0.0",
    docs_url="/docs",
)

@app.on_event("startup")
def on_startup():
    init_db()

app.include_router(auth_router)
app.include_router(users_router)
app.include_router(history_router)
app.include_router(employees_router)
app.include_router(reports_router)
app.include_router(translator_router)
app.include_router(document_format_router)

app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.cors_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── Helper: detect file type ──────────────────────────────────────────────────

ALLOWED_TYPES: dict[str, str] = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}


def detect_file_type(content_type: str, filename: str) -> str:
    if content_type in ALLOWED_TYPES:
        return ALLOWED_TYPES[content_type]
    fn = (filename or "").lower()
    if fn.endswith(".pdf"):
        return "pdf"
    if fn.endswith(".docx"):
        return "docx"
    raise HTTPException(400, "Faqat PDF yoki Word (.docx) fayllar qabul qilinadi")


# ── Helper: save corrected text ────────────────────────────────────────────────

def save_as_pdf(text: str, path: Path) -> None:
    doc = SimpleDocTemplate(
        str(path), pagesize=A4,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
    )
    styles = getSampleStyleSheet()
    style = ParagraphStyle(
        "UzText", parent=styles["Normal"],
        fontSize=12, leading=18, fontName="Helvetica",
    )
    story = []
    for line in text.split("\n"):
        if line.strip():
            story.append(RLParagraph(line.strip(), style))
            story.append(Spacer(1, 6))
    doc.build(story)


def save_as_docx(text: str, path: Path) -> None:
    doc = DocxDocument()
    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())
    doc.save(str(path))


# ── Routes ─────────────────────────────────────────────────────────────────────

@app.get("/health")
def health():
    return {"status": "ok", "model": settings.openai_model}


@app.get("/version")
def version():
    return {
        "version": "2.0.0",
        "model": settings.openai_model,
        "translation_model": settings.translation_model,
        "chunk_size": settings.chunk_size_chars,
        "max_parallel": settings.max_parallel_chunks,
    }


@app.post("/analyze-document", response_model=AnalysisResponse)
async def analyze_document(file: UploadFile = File(...)):
    """
    PDF yoki DOCX faylni yuklang.
    AI barcha chunklarni parallel tahlil qilib to'liq natija qaytaradi.
    """
    start_ms = time.time()

    if not settings.openai_api_key:
        raise HTTPException(500, "OPENAI_API_KEY sozlanmagan")

    # ── 1. Fayl qabul qilish ──────────────────────────────────────────────────
    file_type = detect_file_type(file.content_type or "", file.filename or "")
    content = await file.read()

    max_bytes = settings.max_file_size_mb * 1024 * 1024
    if len(content) > max_bytes:
        raise HTTPException(400, f"Fayl hajmi {settings.max_file_size_mb}MB dan oshmasligi kerak")

    file_id = str(uuid.uuid4())
    save_path = UPLOAD_DIR / f"{file_id}.{file_type}"
    save_path.write_bytes(content)
    logger.info("File saved: %s (%d bytes)", save_path.name, len(content))

    # ── 2. Matn ajratish ──────────────────────────────────────────────────────
    try:
        if file_type == "pdf":
            parse_result = parse_pdf(save_path)
        else:
            parse_result = parse_docx(save_path)
    except Exception as e:
        save_path.unlink(missing_ok=True)
        raise HTTPException(422, f"Fayldan matn o'qib bo'lmadi: {e}")

    if not parse_result.paragraphs:
        save_path.unlink(missing_ok=True)
        raise HTTPException(422, "Faylda o'qib bo'ladigan matn topilmadi")

    paragraphs = parse_result.paragraphs
    full_text = "\n\n".join(p.text for p in paragraphs)
    word_count = len(full_text.split())

    metadata = FileMetadata(
        page_count=parse_result.page_count,
        paragraph_count=len(paragraphs),
        word_count=word_count,
        char_count=len(full_text),
    )
    logger.info("Parsed: %d paragraphs, %d words", len(paragraphs), word_count)

    # ── 3. Chunklarga bo'lish ─────────────────────────────────────────────────
    chunks = build_chunks(paragraphs, chunk_size_chars=settings.chunk_size_chars)
    logger.info("Built %d chunks", len(chunks))

    # ── 4. Parallel tahlil ───────────────────────────────────────────────────
    chunk_results = await analyze_all_chunks(chunks)

    # ── 5. Natijalarni yig'ish ────────────────────────────────────────────────
    corrected_text = "\n\n".join(cr.corrected_text for cr in chunk_results)
    issue_counts, total_issues = build_summary(chunk_results)

    summary = Summary(
        total_chunks=len(chunks),
        total_issues=total_issues,
        issue_counts=issue_counts,
    )

    # ── 6. To'g'irlangan faylni saqlash ──────────────────────────────────────
    fixed_id = str(uuid.uuid4())
    try:
        save_as_pdf(corrected_text, UPLOAD_DIR / f"{fixed_id}.pdf")
        save_as_docx(corrected_text, UPLOAD_DIR / f"{fixed_id}.docx")
    except Exception as e:
        logger.warning("Could not save corrected files: %s", e)

    # fixed_id ni warnings ga qo'shib frontendga uzatamiz
    warnings = list(parse_result.warnings)
    if total_issues == 0:
        warnings.append("Hujjatda aniq korrektura xatosi topilmadi.")

    elapsed_ms = int((time.time() - start_ms) * 1000)
    logger.info("Analysis done: %d issues, %d ms", total_issues, elapsed_ms)

    return AnalysisResponse(
        success=True,
        file_name=file.filename or "unknown",
        file_type=file_type,
        processing_time_ms=elapsed_ms,
        metadata=metadata,
        summary=summary,
        corrected_text=corrected_text,
        chunks=chunk_results,
        warnings=warnings,
        file_id=fixed_id,
    )


@app.get("/download/{file_id}")
def download_file(file_id: str, format: Literal["pdf", "docx"] = "pdf"):
    """To'g'irlangan hujjatni yuklab olish."""
    if not all(c.isalnum() or c == "-" for c in file_id):
        raise HTTPException(400, "Noto'g'ri file_id")

    file_path = UPLOAD_DIR / f"{file_id}.{format}"
    if not file_path.exists():
        raise HTTPException(404, "Fayl topilmadi")

    media_types = {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    return FileResponse(
        path=str(file_path),
        media_type=media_types[format],
        filename=f"tuzatilgan_hujjat.{format}",
    )
